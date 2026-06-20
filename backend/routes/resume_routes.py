from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from io import BytesIO
import shutil, os, uuid, datetime
from dotenv import load_dotenv
from motor.motor_asyncio import AsyncIOMotorClient
from docx import Document
from resume_parser import extract_text_from_pdf, parse_resume_with_llm
from services.recommendation.recommendation_engine import recommend_careers
from services.resume.formatter import format_parsed_profile
from services.resume.generator import generate_resume

load_dotenv()

router = APIRouter(prefix="/api/resume", tags=["resume"])

UPLOAD_DIR = os.path.join(os.getcwd(), "temp_resumes")
os.makedirs(UPLOAD_DIR, exist_ok=True)

MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017")
_db_client = AsyncIOMotorClient(MONGO_URI)
_db = _db_client.career_recommender


@router.post("/upload")
async def upload_resume(file: UploadFile = File(...)):
    """Upload PDF resume, parse with LLM, return recommendations."""
    if not file.filename.endswith(".pdf"):
        raise HTTPException(400, "Only PDF files supported")

    file_id = str(uuid.uuid4())
    path = os.path.join(UPLOAD_DIR, f"{file_id}.pdf")

    try:
        with open(path, "wb") as f:
            shutil.copyfileobj(file.file, f)

        raw_text = extract_text_from_pdf(path)
        if not raw_text.strip():
            raise HTTPException(400, "Could not extract text from PDF")

        parsed = parse_resume_with_llm(raw_text)
        recommendations = recommend_careers(parsed, top_n=5)

        doc = {
            "file_id": file_id,
            "filename": file.filename,
            "parsed_profile": parsed,
            "recommendations": recommendations,
            "timestamp": datetime.datetime.utcnow(),
        }
        await _db.resumes.insert_one(doc)
        os.remove(path)

        return {"id": file_id, "parsed_profile": parsed, "recommendations": recommendations}

    except Exception as e:
        if os.path.exists(path):
            os.remove(path)
        raise HTTPException(500, f"Error: {str(e)}")


@router.post("/generate")
async def generate_resume_from_profile(data: dict):
    """Generate a professional resume text from a structured profile."""
    formatted_profile = format_parsed_profile(data)
    if not formatted_profile["skills"] and not formatted_profile["current_role"]:
        raise HTTPException(400, "Provide at least skills or current role for resume generation")

    try:
        resume_text = generate_resume(formatted_profile)
        return {"resume": resume_text, "profile": formatted_profile}
    except Exception as e:
        raise HTTPException(500, f"Error generating resume: {str(e)}")


@router.post("/generate/word")
async def generate_resume_word_from_profile(data: dict):
    """Generate a downloadable Word document from a structured profile."""
    formatted_profile = format_parsed_profile(data)
    if not formatted_profile["skills"] and not formatted_profile["current_role"]:
        raise HTTPException(400, "Provide at least skills or current role for resume generation")

    try:
        resume_text = generate_resume(formatted_profile)
        buffer = _build_resume_docx_buffer(resume_text)
        headers = {"Content-Disposition": 'attachment; filename="generated_resume.docx"'}
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers,
        )
    except Exception as e:
        raise HTTPException(500, f"Error generating resume word file: {str(e)}")


def _build_resume_docx_buffer(resume_text: str) -> BytesIO:
    document = Document()
    for line in resume_text.splitlines():
        text = line.strip()
        if not text:
            document.add_paragraph("")
            continue
        if text == "---":
            continue
        if text.isupper() and len(text) > 3:
            document.add_heading(text, level=1)
            continue
        document.add_paragraph(text)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
